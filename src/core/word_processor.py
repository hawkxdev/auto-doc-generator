"""
Обработчик Word документов
"""
import re
from pathlib import Path
from docx import Document
from docx.shared import Inches
from src.utils.logger import log_info, log_success, log_warning, log_error


class WordProcessor:
    """Класс для обработки Word документов"""

    def __init__(self, config):
        self.config = config
        self.template_path = None
        self.template_doc = None
        self.placeholders_found = {}

    def load_template(self, template_path: str):
        """Загрузка Word шаблона"""
        self.template_path = Path(template_path)

        if not self.template_path.exists():
            raise FileNotFoundError(f"Word шаблон не найден: {template_path}")

        try:
            self.template_doc = Document(self.template_path)
            self._scan_placeholders()
            log_success(f"Word шаблон загружен: {len(self.placeholders_found)} плейсхолдеров")
            return self.template_doc
        except Exception as e:
            raise RuntimeError(f"Ошибка загрузки Word шаблона: {e}")

    def _scan_placeholders(self):
        """Сканирование плейсхолдеров в шаблоне"""
        self.placeholders_found = {}
        placeholder_pattern = r'\{[^}]+\}'

        for paragraph in self.template_doc.paragraphs:
            placeholders = re.findall(placeholder_pattern, paragraph.text)
            for placeholder in placeholders:
                if placeholder in self.config['placeholders']:
                    self.placeholders_found[placeholder] = self.placeholders_found.get(placeholder, 0) + 1

        for table in self.template_doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    placeholders = re.findall(placeholder_pattern, cell.text)
                    for placeholder in placeholders:
                        if placeholder in self.config['placeholders']:
                            self.placeholders_found[placeholder] = self.placeholders_found.get(placeholder, 0) + 1

    def create_document_from_template(self, row_data, output_path: str):
        """Создание документа из шаблона с заменой плейсхолдеров"""
        if self.template_doc is None:
            raise ValueError("Шаблон не загружен")

        doc = Document(self.template_path)

        text_replacements = 0
        image_insertions = 0
        images_requested = 0

        for placeholder, config_data in self.config['placeholders'].items():
            column_name = config_data['column']
            placeholder_type = config_data['type']

            if column_name not in row_data:
                log_warning(f"Колонка '{column_name}' не найдена в данных")
                continue

            value = row_data[column_name]

            if placeholder_type == 'text':
                replacements = self._replace_text_placeholder(doc, placeholder, str(value))
                text_replacements += replacements
            elif placeholder_type == 'image':
                images_requested += 1
                insertions = self._replace_image_placeholder(doc, placeholder, value)
                image_insertions += insertions

        # NOTE: В полной версии здесь будет обработка колонтитулов
        # self._process_headers_and_footers(doc, row_data)

        doc.save(output_path)
        log_success(f"Документ создан: {text_replacements} замен текста, {image_insertions} изображений")

        return {
            'text_replacements': text_replacements,
            'image_insertions': image_insertions,
            'images_requested': images_requested
        }

    def _replace_text_placeholder(self, doc, placeholder, value):
        """Замена текстового плейсхолдера"""
        replacements = 0

        for paragraph in doc.paragraphs:
            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder, value)
                replacements += paragraph.text.count(value)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if placeholder in cell.text:
                        cell.text = cell.text.replace(placeholder, value)
                        replacements += cell.text.count(value)

        return replacements

    def _replace_image_placeholder(self, doc, placeholder, image_name):
        """Замена плейсхолдера изображением"""
        if not image_name or str(image_name).strip() == '':
            return 0

        image_path = self._find_image_file(image_name)
        if not image_path:
            log_warning(f"Изображение не найдено: {image_name}")
            return 0

        insertions = 0

        for paragraph in doc.paragraphs:
            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder, '')
                try:
                    paragraph.add_run().add_picture(
                        str(image_path),
                        height=Inches(self.config['processing']['image_height_inches'])
                    )
                    insertions += 1
                except Exception as e:
                    log_error(f"Ошибка вставки изображения {image_name}: {e}")

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if placeholder in paragraph.text:
                            paragraph.text = paragraph.text.replace(placeholder, '')
                            try:
                                paragraph.add_run().add_picture(
                                    str(image_path),
                                    height=Inches(self.config['processing']['image_height_inches'])
                                )
                                insertions += 1
                            except Exception as e:
                                log_error(f"Ошибка вставки изображения {image_name}: {e}")

        return insertions

    def _find_image_file(self, image_name):
        """Поиск файла изображения"""
        image_name = str(image_name).strip()

        base_path = Path(self.config['input']['images_folder'])

        for subdirectory in self.config['input']['images_subdirectories'].values():
            subdir_path = base_path / subdirectory
            if not subdir_path.exists():
                continue

            for extension in ['.jpg', '.jpeg', '.png', '.gif', '.bmp']:
                image_path = subdir_path / f"{image_name}{extension}"
                if image_path.exists():
                    return image_path

        return None

    def update_statistics(self, stats, text_replacements, image_insertions, images_requested):
        """Обновление статистики операций"""
        stats.add_text_replacements(text_replacements)
        stats.add_image_insertions(image_insertions)
        stats.add_document_created()

        images_found = image_insertions
        images_not_found = images_requested - image_insertions

        for _ in range(images_found):
            stats.add_image_found()
        for _ in range(images_not_found):
            stats.add_image_not_found()
